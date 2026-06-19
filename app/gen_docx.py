#!/usr/bin/env python3
"""Convert WeatherClock interview notes from Markdown to Word (.docx)"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
import re

doc = Document()

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hf = hs.font
    hf.name = 'Arial'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if i == 1:
        hf.size = Pt(22)
        hf.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    elif i == 2:
        hf.size = Pt(16)
        hf.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    else:
        hf.size = Pt(13)
        hf.color.rgb = RGBColor(0x34, 0x49, 0x5e)

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    shd = p._element.get_or_add_pPr().makeelement(qn('w:shd'), {qn('w:fill'): 'f5f5f5', qn('w:val'): 'clear'})
    p._element.get_or_add_pPr().append(shd)

def add_rich_paragraph(doc, text, bold_keywords=None):
    """Add paragraph with inline `code` and **bold** formatting."""
    p = doc.add_paragraph()
    # Split by inline code first
    parts = re.split(r'(`[^`]+`)', text)
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            run = p.add_run(part.strip('`'))
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
        else:
            # Split by bold markers
            sub = re.split(r'(\*\*[^*]+\*\*)', part)
            for s in sub:
                if s.startswith('**') and s.endswith('**'):
                    run = p.add_run(s.strip('*'))
                    run.font.bold = True
                elif s:
                    run = p.add_run(s)

def add_simple_paragraph(doc, text):
    """Add plain paragraph, stripping bold markers."""
    p = doc.add_paragraph()
    clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    run = p.add_run(clean)

def parse_and_build(doc, md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Code fence
        if line.startswith('```'):
            if in_code:
                add_code_block(doc, '\n'.join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Table detection: collect consecutive table lines
        if '|' in line and line.strip().startswith('|'):
            tbl = [line]
            j = i + 1
            while j < len(lines) and '|' in lines[j] and lines[j].strip().startswith('|'):
                tbl.append(lines[j].rstrip())
                j += 1
            # Filter separator line
            data_lines = [l for l in tbl if not re.match(r'^\|[\s\-:|]+\|$', l.strip())]
            if data_lines:
                headers = [h.strip() for h in data_lines[0].strip('|').split('|')]
                table = doc.add_table(rows=len(data_lines), cols=len(headers))
                table.style = 'Light Grid Accent 1'
                for ri, dl in enumerate(data_lines):
                    cells = [c.strip() for c in dl.strip('|').split('|')]
                    for ci, ct in enumerate(cells):
                        if ci < len(headers):
                            cell = table.rows[ri].cells[ci]
                            cell.text = ct
                            for pp in cell.paragraphs:
                                for rn in pp.runs:
                                    rn.font.size = Pt(9)
                                    if ri == 0:
                                        rn.font.bold = True
                doc.add_paragraph()
            i = j
            continue

        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('---'):
            p = doc.add_paragraph()
            run = p.add_run('─' * 50)
            run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = re.sub(r'^[\s]*[-*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            add_rich_runs_to_para(p, text)
        elif re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            add_rich_runs_to_para(p, text)
        elif line.strip():
            add_rich_paragraph(doc, line.strip())
        else:
            # Blank line spacer
            pass

        i += 1

def add_rich_runs_to_para(p, text):
    """Add runs to an existing paragraph, handling `code` and **bold**."""
    parts = re.split(r'(`[^`]+`)', text)
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            run = p.add_run(part.strip('`'))
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
        else:
            sub = re.split(r'(\*\*[^*]+\*\*)', part)
            for s in sub:
                if s.startswith('**') and s.endswith('**'):
                    run = p.add_run(s.strip('*'))
                    run.font.bold = True
                elif s:
                    run = p.add_run(s)

# Build
md_path = r'C:\Users\l\Desktop\WeatherClock_面试知识点.md'
parse_and_build(doc, md_path)

# Save
out = r'C:\Users\l\Desktop\WeatherClock_面试知识点.docx'
doc.save(out)
print(f'Done: {out}')
